import os
import ast
import json
import docx
import openai
import base64
import logging
import tempfile
import pytesseract
from gtts import gTTS
from io import BytesIO
from flask_cors import CORS
from datetime import datetime
from dotenv import load_dotenv
from pdf2image import convert_from_path
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text
from logging.handlers import RotatingFileHandler
from flask_jwt_extended import JWTManager, create_access_token

# Initialize Flask app
app = Flask(__name__)
CORS(app)

# Load environment variables from .env file
load_dotenv()
app.logger.info(f"Loaded environment variables: {os.environ}")

# Set OpenAI API key
openai.api_key = "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"
app.logger.info(f"OpenAI API key set: {openai.api_key}")

# Configure JWT
app.config["JWT_SECRET_KEY"] = "interviewaxis"
jwt = JWTManager(app)
app.logger.info(f"JWT secret key set: {app.config['JWT_SECRET_KEY']}")

# Set environment variables for external binaries
os.environ['PATH'] += ':/usr/bin:/path/to/poppler/bin'
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
app.logger.info(f"Environment PATH updated: {os.environ['PATH']}")

# Ensure the 'logs' directory exists
logs_dir = "logs"
if not os.path.exists(logs_dir):
    os.makedirs(logs_dir)
    app.logger.info(f"Created logs directory: {logs_dir}")

# Configure logging
log_formatter = logging.Formatter('%(asctime)s %(levelname)s %(message)s')
log_handler = RotatingFileHandler(
    filename=os.path.join(logs_dir, f'logger_{datetime.now().strftime("%Y%m%d%H%M%S")}.log'),
    maxBytes=1 * 1024 * 1024 * 1024,  # 1GB
    backupCount=10  # Keep up to 10 backup logs
)
log_handler.setFormatter(log_formatter)
app.logger.addHandler(log_handler)
app.logger.setLevel(logging.INFO)

# Log application start
app.logger.info("Application started")

def convert_to_json_builder(data, input_text):
    app.logger.info("Starting convert_to_json_builder function")
    def process_section(section_name, instruction):
        app.logger.info(f"Processing section: {section_name}")
        prompt = f"""
        Translate the following resume text into exactly this Python Dict format for the '{section_name}' section :
        {instruction}"""
        prompt += f"""Take exactly what is present in the Summary or Description or Introduction or Objective and return it as a Python Dict unless it is more than a paragraph, in this case, make it more concise. If the text does not contain any field that represents the summary, generate a small paragraph that responds to this purpose""" if section_name == "summary" else f""""""
        prompt += f"""Make sure the name of the of each reference is different than the name of the resume holder. If none are found, leave the items empty.""" if section_name == "references" else f""""""
        prompt += f"""Make sure the skills are different from elements in the other sections, like certifications or languages. If none are found, leave the items empty.""" if section_name == "skills" else f""""""
        prompt += f"""Provide the Python Dict for only the '{section_name}' section without additional commentary keeping complete data integrity especially in experience missions. The output must absolutely be a valid Python Dict :
        Text: {input_text}
        """

        app.logger.info(f"Sending prompt to OpenAI API for section: {section_name}")
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            temperature=0.1,
            max_tokens=4096,
            messages=[
                {"role": "system", "content": "You are a helpful assistant. You are an expert in Python Dict formatting and in reading, making and reviewing resumes."},
                {"role": "user", "content": prompt}
            ]
        )
        app.logger.info(f"Received response from OpenAI API for section: {section_name}")
        app.logger.info(f"OpenAI API usage - Prompt tokens: {response.usage.prompt_tokens}, Completion tokens: {response.usage.completion_tokens}, Total tokens: {response.usage.total_tokens}")

        output = response.choices[0].message['content'].replace("```json", "").replace("```python", "").replace("```", "")

        def clean_nested_json(value):
            if isinstance(value, str):
                try:
                    return ast.literal_eval(value)
                except json.JSONDecodeError:
                    return value
            elif isinstance(value, dict):
                return {k: clean_nested_json(v) for k, v in value.items()}
            elif isinstance(value, list):
                return [clean_nested_json(v) for v in value]
            return value

        cleaned_data = clean_nested_json(output)

        try:
            app.logger.info(f"Successfully processed section: {section_name}")
            return cleaned_data, response.usage.prompt_tokens, response.usage.completion_tokens
        except json.JSONDecodeError:
            app.logger.error(f"Invalid JSON for section '{section_name}'. Please refine the input.")
            return f"Invalid JSON for section '{section_name}'. Please refine the input."

    result = {
        "basics": {
        },
        "sections": {
            "current_skills": {
                "name": "Current Skills",
                "items": [
                    {
                        "name": "",
                        "level": 3,
                        "visible": True
                    }
                ]
            },
            "missing_skills": {
                "name": "Missing Skills",
                "items": [
                    {
                        "name": "",
                        "level": 3,
                        "visible": True
                    }
                ]
            },
            "suggested_missions": {
                "name": "Suggested Missions",
                "items": [
                    ""
                ],
            },
        }
    }

    for section_name, instruction in data.items():
        processed, input_price, output_price = process_section(section_name, instruction)

        if section_name == "basics":
            result["basics"] = processed
        else:
            result["sections"][section_name] = processed
            items = processed.get("items", [])
            if len(items) == 1 and items[0].get("name", "").strip() == "":
                result["sections"][section_name]["items"] = []

    app.logger.info("Finished convert_to_json_builder function")
    return result

def get_missing_skills(resume_text, job_description):
    app.logger.info("Starting get_missing_skills function")
    prompt = f""" I have a resume and a job description. Identify missing skill sets from the job description that are not covered in the resume, and provide multiple sentences for each missing skill that I can directly add to the "Projects" section of my resume.

    Please ensure that each sentence relates to a each skill missing from the job description and the project mentioned in the resume that is missing or not fully covered in the resume. Format the output strictly as JSON so that it can pass json.loads without error.
    so give me the sentences which will create a story of skills missing and should fit in the resume projects . analysis the project from resume and try to create sentence which after adding looks like its a part of project
    The output format should be:
    give always json response in mentioned format I have a resume and a job description. i want the sentence , which i can add to the projects mentioned in the resume .

        Provide the output in the following JSON format and do a detail research and give me data in enhanced form :

        "give me multiple sentences for each skills set  which i can add in the project which are missing give me the points , it should be in json like dictnory answer and should be sentence which i can add directly in the resume manually"
        for example output should be {{"missing_skills": "skill1": ["",""],"skill2": ["".""] }} like a list with all answers which will pass json.loads . dont give additional information   . only json answer dont add ```json also

    Resume: {resume_text}
    Job Description: {job_description}"""

    app.logger.info(f"Sending prompt to OpenAI API for missing skills")
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        temperature=0.3,
        max_tokens=2048,
        messages=[
            {"role": "system", "content": "You are a helpful assistant that provides suggestions for improving resumes based on job descriptions."},
            {"role": "user", "content": prompt}
        ]
    )
    app.logger.info(f"Received response from OpenAI API for missing skills")
    app.logger.info(f"OpenAI API usage - Prompt tokens: {response.usage.prompt_tokens}, Completion tokens: {response.usage.completion_tokens}, Total tokens: {response.usage.total_tokens}")

    app.logger.info("Finished get_missing_skills function")
    return response.choices[0].message['content']

def extract_text_from_docx(file_path):
    app.logger.info("Starting extract_text_from_docx function")
    text = ""
    try:
        doc = docx.Document(file_path)

        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append(para.text)

        tables = []
        for table in doc.tables:
            table_text = "\n[Table Start]\n"
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + "\n"
            table_text += "[Table End]\n"
            tables.append(table_text)

        text += "\n".join(paragraphs) + "\n"

        text += "\n".join(tables)

    except Exception as e:
        app.logger.error(f"Error extracting text from DOCX: {e}")
    app.logger.info("Finished extract_text_from_docx function")
    return text

def extract_text_from_pdf(file_path):
    app.logger.info("Starting extract_text_from_pdf function")
    try:
        text = extract_text(file_path)
        app.logger.info("Finished extract_text_from_pdf function")
        return text
    except Exception as e:
        app.logger.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_image_pdf(file_path):
    app.logger.info("Starting extract_text_from_image_pdf function")
    text = ""
    try:
        images = convert_from_path(file_path, first_page=1, last_page=1)
        for image in images:
            text += pytesseract.image_to_string(image)
    except Exception as e:
        app.logger.error(f"Error extracting text from image-based PDF: {e}")
    app.logger.info("Finished extract_text_from_image_pdf function")
    return text

def extract_full_text_from_file(file_path):
    app.logger.info("Starting extract_full_text_from_file function")
    file_extension = file_path.lower().split('.')[-1]
    app.logger.info(f"File extension: {file_extension}")
    if file_extension == 'pdf':
        text = extract_text_from_pdf(file_path)
        if not text.strip():
            app.logger.info("No text found in PDF, attempting OCR extraction.")
            text = extract_text_from_image_pdf(file_path)
    elif file_extension == 'docx':
        text = extract_text_from_docx(file_path)
    else:
        app.logger.error("Unsupported file type.")
        return "Unsupported file type."
    app.logger.info("Finished extract_full_text_from_file function")
    return text

def text_to_speech(text):
    app.logger.info("Starting text_to_speech function")
    if text:
        tts = gTTS(text=text, slow=False)
        audio_buffer = BytesIO()
        tts.write_to_fp(audio_buffer)
        audio_data = audio_buffer.getvalue()
        app.logger.info("Successfully converted text to speech")
        return audio_data
    else:
        app.logger.error("No text provided for speech synthesis")
        return None

@app.route('/modifier/resume_builder', methods=['POST'])
def resume_builder():
    app.logger.info("Received request for /modifier/resume_builder")
    if 'file' not in request.files:
        app.logger.error("No file provided in request")
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        app.logger.error("No selected file")
        return jsonify({'error': 'No selected file'}), 400
    
    data = {
        "basics": {
            "name": "",
            "email": "",
            "phone": "",
            "headline": "",
            "location": "",
            "url": {
                "href": "",
                "label": ""
            },
            "picture": {
                "url": ""
            },
            "customFields": []
        },
        "skills": {
            "name": "Skills",
            "items": [
                {
                    "name": "",
                    "level": 3,
                    "visible": True
                }
            ]
        },
        "summary": {
            "name": "Summary",
            "content": "",
            "visible": True
        },
        "profiles": {
            "name": "Profiles",
            "items": [
                {
                    "name": "",
                    "url": {
                        "href": "",
                        "label": ""
                    }
                }
            ]
        },
        "projects": {
            "name": "Projects",
            "items": [
                {
                    "name": "",
                    "description": "",
                    "skills": [
                        ""
                    ]
                }
            ]
        },
        "interests": {
            "name": "Interests",
            "items": [
                {
                    "name": ""
                }
            ]
        },
        "languages": {
            "name": "Languages",
            "items": [
                {
                    "name": "",
                    "level": 0,
                    "proficiency": "100%"
                }
            ]
        },
        "volunteer": {
            "name": "Volunteering",
            "items": [
                {
                    "position": "",
                    "company": "",
                    "location": "",
                    "date": "",
                    "summary": "",
                    "visible": True
                }
            ]
        },
        "experience": {
            "name": "Experience",
            "items": [
                {
                    "position": "",
                    "company": "",
                    "location": "",
                    "date": "",
                    "summary": "",
                    "missions": [
                        ""
                    ],
                    "visible": True
                }
            ]
        },
        "references": {
            "name": "References",
            "items": [],
            "columns": 1,
            "visible": True,
            "separateLinks": True
        },
        "publications": {
            "name": "Publications",
            "items": [
                {
                    "name": "",
                    "description": "",
                    "issuer": "",
                    "date": "",
                    "url": {
                        "href": "",
                        "label": ""
                    }
                }
            ]
        },
        "certifications": {
            "name": "Certifications",
            "items": [
                {
                    "name": "",
                    "description": "",
                    "issuer": "",
                    "date": "",
                    "url": {
                        "href": "",
                        "label": ""
                    }
                }
            ]
        },
        "awards": {
            "name": "Awards",
            "items": [
                {
                    "name": "",
                    "description": "",
                    "issuer": "",
                    "date": ""
                }
            ]
        },
        "education": {
            "name": "Education",
            "items": [
                {
                    "institution": "",
                    "studyType": "",
                    "area": "",
                    "date": "",
                    "summary": "",
                    "score": "",
                    "url": {
                        "href": "",
                        "label": ""
                    }
                }
            ]
        }
    }

    job_description = request.form.get("job_description")

    app.logger.info(f"Processing file: {file.filename}")
    file_path = os.path.join(tempfile.gettempdir(), secure_filename(file.filename))
    file.save(file_path)
    app.logger.info(f"File saved to temporary path: {file_path}")

    extracted_text = extract_full_text_from_file(file_path)
    app.logger.info(f"Extracted text from file: {file.filename}")

    resume_json = convert_to_json_builder(data, extracted_text)
    app.logger.info(f"Generated JSON from extracted text for file: {file.filename}")

    skills_to_add = list(resume_json["sections"]["skills"]["items"])
    missing_skills = json.loads(get_missing_skills(extracted_text, job_description))["missing_skills"]
    app.logger.info(f"Identified missing skills: {missing_skills}")

    missing_skills_to_add = []
    suggested_missions_to_add = []
    for skill in missing_skills:
        skills_to_add.append({
            "name": skill,
            "level": 3,
            "description": "",
            "visible": True
        })
        missing_skills_to_add.append({
            "name": skill,
            "level": 3,
            "description": "",
            "visible": True
        })
        for item in missing_skills[skill]:
            suggested_missions_to_add.append(item)

    resume_json["sections"]["current_skills"]["items"] = resume_json["sections"]["skills"]["items"]
    resume_json["sections"]["skills"]["items"] = skills_to_add
    resume_json["sections"]["missing_skills"]["items"] = missing_skills_to_add
    resume_json["sections"]["suggested_missions"]["items"] = suggested_missions_to_add

    app.logger.info("Finished processing /modifier/resume_builder request")
    return jsonify(resume_json), 200

@app.route('/modifier/get_token', methods=['POST'])
def get_token():
    app.logger.info("Received request for /modifier/get_token")
    user_id = request.form.get("user_id")
    user_email = request.form.get("user_email")
    app.logger.info(f"Generating token for user_id: {user_id}, user_email: {user_email}")
    access_token = create_access_token(identity={"user_id": user_id, "user_email": user_email})
    app.logger.info("Finished processing /modifier/get_token request")
    return jsonify({"access_token": access_token}), 200

@app.route('/transform/upload_file', methods=['POST'])
def upload_file():
    app.logger.info("Received request for /transform/upload_file")
    if 'file' not in request.files:
        app.logger.error("No file provided in request")
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        app.logger.error("No selected file")
        return jsonify({'error': 'No selected file'}), 400
    
    data = {
        "basics": {
            "name": "",
            "headline": "",
            "email": "",
            "phone": "",
            "location": "",
            "url": {
                "label": "",
                "href": "https://default"
            },
            "customFields": [],
            "picture": {
                "url": "",
                "size": 64,
                "aspectRatio": 1,
                "borderRadius": 0,
                "effects": {
                    "hidden": False,
                    "border": False,
                    "grayscale": False,
                },
            },
        },
        "skills": {
            "id": "skills",
            "name": "Skills",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "name": "",
                    "level": 1,
                    "keywords": [],
                    "description": "",
                    "visible": True
                }
            ]
        },
        "summary": {
            "id": "summary",
            "name": "Summary",
            "content": "",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
        },
        "profiles": {
            "id": "profiles",
            "name": "Profiles",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "icon": "",
                    "network": "",
                    "username": "",
                    "visible": True
                }
            ]
        },
        "projects": {
            "id": "projects",
            "name": "Projects",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "date": "",
                    "name": "",
                    "summary": "",
                    "keywords": [],
                    "description": "",
                    "visible": True
                }
            ]
        },
        "interests": {
            "name": "Interests",
            "id": "interests",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "name": "",
                    "keywords": [],
                    "visible": True
                }
            ]
        },
        "languages": {
            "name": "Languages",
            "id": "languages",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "name": "",
                    "level": 0,
                    "description": "",
                    "visible": True
                }
            ]
        },
        "volunteer": {
            "name": "Volunteering",
            "id": "volunteer",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "date": "",
                    "summary": "",
                    "location": "",
                    "position": "",
                    "organization": "",
                    "visible": True
                }
            ]
        },
        "experience": {
            "name": "Experience",
            "id": "experience",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "date": "",
                    "company": "",
                    "summary": "",
                    "location": "",
                    "position": "",
                    "visible": True
                }
            ]
        },
        "references": {
            "name": "References",
            "id": "references",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "name": "",
                    "summary": "",
                    "description": "",
                    "visible": True
                }
            ],
            "visible": True,
            "separateLinks": True
        },
        "publications": {
            "name": "Publications",
            "id": "publications",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "date": "",
                    "name": "",
                    "summary": "",
                    "publisher": "",
                    "visible": True
                }
            ]
        },
        "certifications": {
            "id": "certifications",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "name": "Certifications",
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "date": "",
                    "name": "",
                    "issuer": "",
                    "summary": "",
                    "visible": True
                }
            ]
        },
        "awards": {
            "id": "awards",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "name": "Awards",
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "date": "",
                    "title": "",
                    "awarder": "",
                    "summary": "",
                    "visible": True
                }
            ]
        },
        "education": {
            "id": "education",
            "columns": 1,
            "separateLinks": True,
            "visible": True,
            "name": "Education",
            "items": [
                {
                    "url": {
                        "href": "",
                        "label": ""
                    },
                    "area": "",
                    "date": "",
                    "score": "",
                    "summary": "",
                    "studyType": "",
                    "institution": "",
                    "visible": True
                }
            ]
        }        
    }

    job_description = request.form.get("job_description")

    app.logger.info(f"Processing file: {file.filename}")
    file_path = os.path.join(tempfile.gettempdir(), secure_filename(file.filename))
    file.save(file_path)
    app.logger.info(f"File saved to temporary path: {file_path}")

    extracted_text = extract_full_text_from_file(file_path)
    app.logger.info(f"Extracted text from file: {file.filename}")

    resume_json = convert_to_json_builder(data, extracted_text)
    app.logger.info(f"Generated JSON from extracted text for file: {file.filename}")

    skills_to_add = list(resume_json["sections"]["skills"]["items"])
    missing_skills = json.loads(get_missing_skills(extracted_text, job_description))["missing_skills"]
    app.logger.info(f"Identified missing skills: {missing_skills}")

    suggested_missions_to_add = []
    for skill in missing_skills:
        skills_to_add.append({
            "name": skill,
            "level": 3,
            "keywords": [],
            "description": "",
            "visible": True
        })
        for item in missing_skills[skill]:
            suggested_missions_to_add.append({
                "url": {
                    "href": "",
                    "label": ""
                },
                "date": "",
                "company": "",
                "summary": item,
                "location": "",
                "position": skill,
                "visible": True
            })

    resume_json["sections"]["skills"]["items"] = skills_to_add
    resume_json["sections"]["custom"] = {}
    resume_json["sections"]["experience"]["items"] += suggested_missions_to_add

    app.logger.info("Finished processing /transform/upload_file request")
    return jsonify(resume_json), 200

@app.route("/interviewbot/question", methods=["POST"])
def get_question():
    app.logger.info("Received request for /interviewbot/question")
    if request.method == "POST":
        data = request.json
        sub_category = data["sub_category"]
        level_type = data["level_type"]

        app.logger.info(f"Generating questions for sub-category: {sub_category}, level: {level_type}")
        try:
            prompt = f"""Generate 20 interview questions on the following topics : {sub_category}

                The difficulty level of the questions must be : {level_type}
                
                The questions need to be very concise (1 sentence each).

                The questions must not be repeated.
        
                Your output should be a JSON Array of questions.
            """

            response = openai.ChatCompletion.create(
                model="gpt-4o-mini",
                temperature=0.3,
                max_tokens=2048,
                messages=[{"role": "user", "content": prompt}],
            )
            app.logger.info(f"Received response from OpenAI API for questions")
            app.logger.info(f"OpenAI API usage - Prompt tokens: {response.usage.prompt_tokens}, Completion tokens: {response.usage.completion_tokens}, Total tokens: {response.usage.total_tokens}")

            questions = ast.literal_eval(
                response.choices[0]
                .message["content"]
                .replace("```json", "")
                .replace("```python", "")
                .replace("```", "")
            )

            data = []

            for question in questions:
                audio = text_to_speech(question)
                if audio:
                    audio_base64 = base64.b64encode(audio).decode("utf-8")
                else:
                    audio_base64 = ""
                data.append({"question": question, "audio": audio_base64})

            app.logger.info(f"Generated {len(questions)} questions for sub-category: {sub_category}, level: {level_type}")
            return jsonify(data)
        except Exception as e:
            app.logger.error(f"Error generating questions: {e}")
            return jsonify({"error": "No questions could be generated for the selected sub-category and level type."}), 500

@app.route("/interviewbot/answer", methods=["POST"])
def check_answer():
    app.logger.info("Received request for /interviewbot/answer")
    data = request.json
    user_answer = data["user_answer"]

    try:
        question = data["question"]

        prompt = f"""You are a highly knowledgeable expert in the topic of the question, tasked with evaluating a candidate's response during an interview. Analyze the answer carefully to determine if it addresses the question in a technically accurate, relevant, and original manner.

            Input:

                - Candidate's Answer: {user_answer}
                - Question: {question}

            Output: Return a JSON object with the following fields:

                - score: Rate the candidate's answer using the scale below:
                    - 0–49: The answer is incorrect, insufficient, irrelevant, empty repeats the question, or contains major conceptual errors and inaccuracies.
                    - 50–100: The answer is correct, technically sufficient, and relevant to the question.
                - explanation: Provide feedback to the candidate:
                    - If the answer is incorrect, incomplete, or irrelevant, explain why and encourage the candidate to focus on providing specific, technical details.
                    - If the answer repeats the question, explicitly mention that repeating the question does not demonstrate understanding and a valid response must include original content.
                    - If the answer is conceptually incorrect (e.g., stating that deep copies affect the original object), correct these misunderstandings with clear and concise explanations.
                    - If the answer is correct, acknowledge its strengths and suggest ways to make it even clearer or more comprehensive.

            Pay special attention to:

                - Answers that are identical or similar to the question: Assign a score of 0 and explain that a valid response must contain original content.
                - Empty or irrelevant answers: Assign a score of 0 and encourage the candidate to attempt an answer.
                - Ensure your explanation is concise (2 sentences) and written in a professional and supportive tone.
                - Write the explanation as if addressing the candidate directly without mentioning the expected answer.
                - Use an encouraging tone and praise the effort of the candidate.
        """

        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            temperature=1,
            max_tokens=512,
            messages=[{"role": "user", "content": prompt}],
        )
        app.logger.info(f"Received response from OpenAI API for answer evaluation")
        app.logger.info(f"OpenAI API usage - Prompt tokens: {response.usage.prompt_tokens}, Completion tokens: {response.usage.completion_tokens}, Total tokens: {response.usage.total_tokens}")

        parsed_response = ast.literal_eval(
            response.choices[0]
            .message["content"]
            .replace("```json", "")
            .replace("```python", "")
            .replace("```", "")
            .replace("true", "True")
            .replace("false", "False")
        )

        app.logger.info(f"Evaluated answer for question: {question}")
        return jsonify(
            {
                "user_answer": user_answer,
                "question": question,
                "similarity_score": parsed_response["score"],
                "explanation": parsed_response["explanation"],
            }
        )
    except Exception as e:
        app.logger.error(f"Error evaluating answer: {e}")
        return jsonify({"error": "Failed to evaluate the answer."}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port='5050', debug=True)